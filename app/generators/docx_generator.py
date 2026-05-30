"""Word 文档生成器 — 导出协议文档"""
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.models.protocol import (
    Project, Device, DeviceInterface, Protocol, StatusVariable, Connection, BusConfig,
    UARTFrameConfig, CANFrameConfig, EthernetFrameConfig,
    EthernetParams, RS422Params, RS232Params, CANParams, CANFDParams,
)
from app.models.enums import *

CATEGORY_NAMES = {
    ProtocolCategory.CONTROL_REQUEST: "控制/请求指令",
    ProtocolCategory.PERIODIC_REPORT: "周期上报消息",
    ProtocolCategory.STATUS_CHANGE: "状态变化上报消息",
    ProtocolCategory.EXECUTION_FEEDBACK: "执行结果反馈消息",
}

DTYPE_SIZES = {
    DataType.UINT8: 1, DataType.UINT16: 2, DataType.UINT32: 4, DataType.UINT64: 8,
    DataType.INT8: 1, DataType.INT16: 2, DataType.INT32: 4, DataType.INT64: 8,
    DataType.FLOAT: 4, DataType.DOUBLE: 8, DataType.BOOL: 1, DataType.STRING: 0,
}

COMM_METHODS = {
    "control_request": "请求-应答",
    "periodic_report": "周期发送",
    "status_change": "变化触发",
    "execution_feedback": "应答返回",
}


def _set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def _add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    return table


def generate_docx(project: Project, output_path: str):
    doc = Document()

    # ── Default font ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)

    # ── Cover ──
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project.name or "通讯协议文档")
    run.bold = True
    run.font.size = Pt(28)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\n版本: {project.version}").font.size = Pt(12)
    info.add_run(f"\n作者: {project.author or '-'}").font.size = Pt(12)
    info.add_run(f"\n日期: {datetime.now().strftime('%Y-%m-%d')}").font.size = Pt(12)
    info.add_run(f"\n全局字节序: {'大端 (Big-Endian)' if project.endian == Endian.BIG else '小端 (Little-Endian)'}").font.size = Pt(12)

    doc.add_page_break()

    # ── 1. Overview ──
    doc.add_heading('一、工程概述', level=1)
    _add_styled_table(doc,
        ["项目", "内容"],
        [
            ["工程名称", project.name],
            ["版本", project.version],
            ["作者", project.author],
            ["全局字节序", "大端" if project.endian == Endian.BIG else "小端"],
            ["设备数量", str(len(project.devices))],
            ["总线数量", str(len(project.bus_configs))],
            ["连线数量", str(len(project.connections))],
        ]
    )

    if project.description:
        doc.add_paragraph(f"描述: {project.description}")

    # ── 2. Bus Configurations ──
    doc.add_heading('二、总线配置', level=1)
    for bc in project.bus_configs:
        doc.add_heading(f'2.{project.bus_configs.index(bc) + 1} {bc.name} ({bc.type.value.upper()})', level=2)
        _add_styled_table(doc, ["属性", "值"], [
            ["总线名称", bc.name],
            ["总线类型", bc.type.value.upper()],
        ])

        # Frame config detail
        cfg = bc.frame_config
        doc.add_heading('帧格式参数', level=3)
        if isinstance(cfg, UARTFrameConfig):
            start_mode_text = "固定值" if cfg.start_flag_mode == "fixed" else "独立配置"
            end_mode_text = "固定值" if cfg.end_flag_mode == "fixed" else "独立配置"
            _add_styled_table(doc, ["参数", "值"], [
                ["起始标志长度", f"{cfg.start_flag_len} 字节" if cfg.start_flag_len > 0 else "不使用"],
                ["起始标志模式", start_mode_text],
                ["起始标志固定值", cfg.start_flag_value or "-"],
                ["信息标识长度", f"{cfg.msg_id_len} 字节"],
                ["信息长度字段长", f"{cfg.frame_len_len} 字节"],
                ["信息内容最大长度", f"{cfg.data_max_len} 字节"],
                ["CRC校验长度", f"{cfg.crc_len} 字节" if cfg.crc_len > 0 else "不使用"],
                ["CRC校验类型", cfg.crc_type.value],
                ["结束标志长度", f"{cfg.end_flag_len} 字节" if cfg.end_flag_len > 0 else "不使用"],
                ["结束标志模式", end_mode_text],
                ["结束标志固定值", cfg.end_flag_value or "-"],
                ["字节序", "大端" if cfg.endian == Endian.BIG else "小端"],
            ])
        elif isinstance(cfg, CANFrameConfig):
            _add_styled_table(doc, ["参数", "值"], [
                ["仲裁域ID", cfg.arbitration_id],
                ["帧类型", cfg.frame_type.value],
                ["DLC", str(cfg.dlc)],
                ["BRS", "启用" if cfg.brs else "未启用"],
            ])
        elif isinstance(cfg, EthernetFrameConfig):
            _add_styled_table(doc, ["参数", "值"], [
                ["协议", cfg.protocol],
                ["头部长度", f"{cfg.header_len} 字节"],
                ["消息类型偏移", str(cfg.msg_type_offset)],
                ["消息类型长度", f"{cfg.msg_type_len} 字节"],
                ["序列号偏移", str(cfg.seq_offset)],
                ["数据长度偏移", str(cfg.data_len_offset)],
                ["数据区偏移", str(cfg.data_offset)],
                ["数据区最大长度", f"{cfg.data_max_len} 字节"],
                ["CRC校验类型", cfg.crc_type.value],
                ["校验字偏移", str(cfg.crc_offset)],
            ])

    # ── 3. Network Topology ──
    doc.add_heading('三、网络拓扑', level=1)
    if project.connections:
        rows = []
        for conn in project.connections:
            from_name = ""
            to_name = ""
            for d in project.devices:
                if d.id == conn.from_device_id:
                    from_name = d.name
                    for i in d.interfaces:
                        if i.id == conn.from_interface_id:
                            from_name += f" ({i.name})"
                if d.id == conn.to_device_id:
                    to_name = d.name
                    for i in d.interfaces:
                        if i.id == conn.to_interface_id:
                            to_name += f" ({i.name})"
            rows.append([from_name, "->", to_name, conn.label])
        _add_styled_table(doc,
            ["源设备 (接口)", "方向", "目标设备 (接口)", "标注"],
            rows,
        )

    for device in project.devices:
        doc.add_heading(f'3.{project.devices.index(device) + 1} 设备: {device.name}', level=2)
        _add_styled_table(doc,
            ["属性", "内容"],
            [
                ["设备名称", device.name],
                ["描述", device.description or "-"],
                ["接口数量", str(len(device.interfaces))],
                ["状态量数量", str(len(device.status_variables))],
            ]
        )

    # ── 4. Status Variables ──
    doc.add_heading('四、状态量定义', level=1)
    for device in project.devices:
        if not device.status_variables:
            continue
        doc.add_heading(f'4.{project.devices.index(device) + 1} {device.name}', level=2)

        rows = []
        for sv in device.status_variables:
            ref_count = project.status_var_ref_count(sv.id)
            rows.append([
                sv.name, sv.data_type.value, str(sv.byte_length),
                sv.unit, sv.meaning, sv.remarks, str(ref_count),
            ])
        _add_styled_table(doc,
            ["名称", "数据类型", "字节长度", "单位", "含义", "备注", "引用数"],
            rows,
        )

    # ── 5. Communication Interfaces ──
    doc.add_heading('五、通讯接口', level=1)
    for device in project.devices:
        if not device.interfaces:
            continue
        doc.add_heading(f'5.{project.devices.index(device) + 1} {device.name}', level=2)

        for iface in device.interfaces:
            bc = project.find_bus_config(iface.bus_config_id)
            bus_type_name = bc.type.value.upper() if bc else "?"
            bus_name = bc.name if bc else "?"
            doc.add_heading(f'接口: {iface.name} (绑定总线: {bus_name} [{bus_type_name}])', level=3)

            rows = [
                ["接口名称", iface.name],
                ["绑定总线", f"{bus_name} ({bus_type_name})"],
            ]
            _add_styled_table(doc, ["参数", "值"], rows)

    # ── 6. Protocol Definitions ──
    doc.add_heading('六、通讯协议定义', level=1)
    for device in project.devices:
        for iface in device.interfaces:
            if not iface.protocols:
                continue
            bc = project.find_bus_config(iface.bus_config_id)
            bc_name = bc.name if bc else "?"
            bus_type_name = bc.type.value.upper() if bc else "?"
            doc.add_heading(f'6.{project.devices.index(device) + 1}.{device.interfaces.index(iface) + 1} {device.name} - {iface.name} ({bc_name} [{bus_type_name}])', level=2)

            for proto in iface.protocols:
                cat_name = CATEGORY_NAMES.get(proto.category, proto.category.value)
                doc.add_heading(f'协议: {proto.name} ({cat_name})', level=3)

                _add_styled_table(doc,
                    ["属性", "值"],
                    [
                        ["协议名称", proto.name],
                        ["分类", cat_name],
                        ["通讯方式", proto.comm_method],
                        ["上报周期", f"{proto.report_period_ms} ms" if proto.category == ProtocolCategory.PERIODIC_REPORT else "-"],
                        ["变化阈值", proto.change_threshold or "-"],
                    ]
                )

                cfg = proto.frame_config
                doc.add_heading('帧格式参数', level=4)

                if isinstance(cfg, UARTFrameConfig):
                    start_mode_text = "固定值" if cfg.start_flag_mode == "fixed" else "独立配置"
                    end_mode_text = "固定值" if cfg.end_flag_mode == "fixed" else "独立配置"
                    _add_styled_table(doc, ["参数", "值"], [
                        ["起始标志长度", f"{cfg.start_flag_len} 字节" if cfg.start_flag_len > 0 else "不使用"],
                        ["起始标志模式", start_mode_text],
                        ["起始标志固定值", cfg.start_flag_value or "-"],
                        ["信息标识长度", f"{cfg.msg_id_len} 字节"],
                        ["信息长度字段长", f"{cfg.frame_len_len} 字节"],
                        ["信息内容最大长度", f"{cfg.data_max_len} 字节"],
                        ["CRC校验长度", f"{cfg.crc_len} 字节" if cfg.crc_len > 0 else "不使用"],
                        ["CRC校验类型", cfg.crc_type.value],
                        ["结束标志长度", f"{cfg.end_flag_len} 字节" if cfg.end_flag_len > 0 else "不使用"],
                        ["结束标志模式", end_mode_text],
                        ["结束标志固定值", cfg.end_flag_value or "-"],
                        ["字节序", "大端" if cfg.endian == Endian.BIG else "小端"],
                    ])
                    if proto.start_flag_value:
                        doc.add_paragraph(f"协议级起始标志值: {proto.start_flag_value}")
                    if proto.end_flag_value:
                        doc.add_paragraph(f"协议级结束标志值: {proto.end_flag_value}")
                elif isinstance(cfg, CANFrameConfig):
                    _add_styled_table(doc, ["参数", "值"], [
                        ["仲裁域ID", cfg.arbitration_id],
                        ["帧类型", cfg.frame_type.value],
                        ["DLC", str(cfg.dlc)],
                        ["BRS", "启用" if cfg.brs else "未启用"],
                    ])
                elif isinstance(cfg, EthernetFrameConfig):
                    _add_styled_table(doc, ["参数", "值"], [
                        ["协议", cfg.protocol],
                        ["头部长度", f"{cfg.header_len} 字节"],
                        ["消息类型偏移", str(cfg.msg_type_offset)],
                        ["消息类型长度", f"{cfg.msg_type_len} 字节"],
                        ["序列号偏移", str(cfg.seq_offset)],
                        ["数据长度偏移", str(cfg.data_len_offset)],
                        ["数据区偏移", str(cfg.data_offset)],
                        ["数据区最大长度", f"{cfg.data_max_len} 字节"],
                        ["CRC校验类型", cfg.crc_type.value],
                        ["校验字偏移", str(cfg.crc_offset)],
                    ])

                if proto.fields:
                    doc.add_heading('字段定义', level=4)
                    field_rows = []
                    offset = 0
                    for fld in proto.fields:
                        end = offset + fld.byte_length - 1
                        byte_label = f"{offset}-{end}" if fld.byte_length > 1 else str(offset)
                        field_rows.append([
                            byte_label,
                            fld.name,
                            str(fld.byte_length),
                            fld.data_type.value,
                            fld.description,
                            fld.unit or "-",
                        ])
                        offset += fld.byte_length
                    _add_styled_table(doc,
                        ["字节号", "名称", "字节数", "数据类型", "含义", "单位"],
                        field_rows,
                    )

    # ── 7. Appendix ──
    doc.add_heading('七、附录: CRC 校验算法说明', level=1)
    doc.add_paragraph("本协议中可能使用的校验算法如下:")
    crc_info = [
        ("CRC-8", "多项式 0x07, 初始值 0x00, 8位", "短帧快速校验"),
        ("CRC-16/CCITT", "多项式 0x1021, 初始值 0xFFFF, 16位", "XMODEM, 通用数据通讯"),
        ("CRC-16/MODBUS", "多项式 0x8005, 初始值 0xFFFF, 16位, 低位在前", "Modbus RTU 协议"),
        ("CRC-32", "多项式 0xEDB88320, 初始值 0xFFFFFFFF, 32位", "Ethernet, ZIP, PNG"),
        ("SUM-8", "字节累加, 取低8位", "简单校验和"),
        ("XOR-8", "字节异或, 8位", "快速异或校验"),
    ]
    _add_styled_table(doc, ["算法", "参数", "典型应用"], crc_info)

    # ── Save ──
    doc.save(output_path)
